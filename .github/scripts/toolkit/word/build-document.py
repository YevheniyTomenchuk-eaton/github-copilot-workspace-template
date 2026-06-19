#!/usr/bin/env python3
"""Build a polished, on-brand Word .docx document from a JSON content spec.

Usage:
    python build-document.py <spec.json> <output.docx>

Renders a cover page (colored title block, subtitle, author, date, accent
rule), a styled body with brand-colored headings, an optional lead paragraph,
callout boxes, bullet lists, banded tables, and a page footer with the
document title and page number.

Spec shape (see the office-documents skill for the full reference):
    {
      "title": "Vendor Evaluation Report",
      "subtitle": "Procurement - Confidential",
      "author": "Dana Ruiz",
      "date": "2026-06-19",
      "sections": [
        {
          "heading": "Summary",
          "lead": "Three vendors were assessed.",
          "paragraphs": ["Vendor B scored highest."],
          "callout": "Recommendation: proceed with Vendor B.",
          "bullets": ["Cost: 30%", "Security: 40%"],
          "table": {"headers": ["Vendor", "Score"], "rows": [["A", "78"]]}
        }
      ]
    }

Emits machine-readable KEY=value lines on success:
    OUTPUT=<absolute path to the .docx>
    SECTIONS=<number of sections>

Requires: python-docx  (pip install python-docx)
"""

import json
import os
import sys

PRIMARY = "1F4E79"
ACCENT = "00B3A4"
INK = "1A2233"
MUTED = "5B6B7F"
BAND = "EEF3F9"
WHITE = "FFFFFF"


def _shade(element, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_color)
    element.append(shading)


def _rgb(hex_color):
    from docx.shared import RGBColor

    return RGBColor.from_string(hex_color)


def _set_cell_background(cell, hex_color):
    _shade(cell._tc.get_or_add_tcPr(), hex_color)


def _paragraph_shading(paragraph, hex_color):
    _shade(paragraph._p.get_or_add_pPr(), hex_color)


def _add_field(paragraph, field):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _cover(document, spec):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(72)

    title = document.add_paragraph()
    _paragraph_shading(title, PRIMARY)
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("  " + spec.get("title", "Untitled"))
    run.font.size = Pt(34)
    run.font.bold = True
    run.font.color.rgb = _rgb(WHITE)
    run.font.name = "Calibri"

    if spec.get("subtitle"):
        subtitle = document.add_paragraph()
        sub_run = subtitle.add_run(spec["subtitle"])
        sub_run.font.size = Pt(16)
        sub_run.font.color.rgb = _rgb(MUTED)
        sub_run.font.name = "Calibri"

    rule = document.add_paragraph()
    rule_run = rule.add_run("\u2014" * 18)
    rule_run.font.color.rgb = _rgb(ACCENT)

    meta_bits = [bit for bit in (spec.get("author"), spec.get("date")) if bit]
    if meta_bits:
        meta = document.add_paragraph()
        meta_run = meta.add_run("   |   ".join(meta_bits))
        meta_run.font.size = Pt(11)
        meta_run.font.color.rgb = _rgb(MUTED)
        meta_run.font.name = "Calibri"

    document.add_page_break()


def _style_base(document):
    from docx.shared import Pt

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = _rgb(INK)

    for level, size in ((1, 16), (2, 13)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(PRIMARY)


def _footer(document, text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text + "    |    ")
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb(MUTED)
    _add_field(paragraph, "PAGE")


def _callout(document, text):
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    _paragraph_shading(paragraph, BAND)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Pt(10)
    paragraph.paragraph_format.right_indent = Pt(10)
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = _rgb(PRIMARY)
    run.font.name = "Calibri"


def _table(document, table_spec):
    from docx.shared import Pt

    headers = table_spec.get("headers", [])
    rows = table_spec.get("rows", [])
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for column, label in enumerate(headers):
        cell = table.rows[0].cells[column]
        _set_cell_background(cell, PRIMARY)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(str(label))
        run.font.bold = True
        run.font.color.rgb = _rgb(WHITE)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for column, value in enumerate(row):
            if column < len(cells):
                if row_index % 2 == 1:
                    _set_cell_background(cells[column], BAND)
                run = cells[column].paragraphs[0].add_run(str(value))
                run.font.size = Pt(10)
                run.font.name = "Calibri"


def build(spec_path, output_path):
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("ERROR=python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    with open(spec_path, encoding="utf-8") as spec_file:
        spec = json.load(spec_file)

    document = Document()
    _style_base(document)
    _cover(document, spec)
    _footer(document, spec.get("title", ""))

    sections = spec.get("sections", [])
    for section in sections:
        if section.get("heading"):
            document.add_heading(section["heading"], level=1)
        if section.get("lead"):
            lead = document.add_paragraph()
            lead_run = lead.add_run(section["lead"])
            lead_run.font.size = Pt(13)
            lead_run.font.color.rgb = _rgb(MUTED)
            lead_run.font.name = "Calibri"
        for paragraph in section.get("paragraphs", []):
            document.add_paragraph(str(paragraph))
        if section.get("callout"):
            _callout(document, section["callout"])
        for bullet in section.get("bullets", []):
            document.add_paragraph(str(bullet), style="List Bullet")
        if section.get("table"):
            _table(document, section["table"])

    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document.save(output_path)

    print(f"OUTPUT={output_path}")
    print(f"SECTIONS={len(sections)}")


def main():
    if len(sys.argv) != 3:
        print("ERROR=usage: python build-document.py <spec.json> <output.docx>")
        sys.exit(1)
    build(sys.argv[1], sys.argv[2])


if __name__ == "__main__":
    main()
